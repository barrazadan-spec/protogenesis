from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(29, 111, 120)
ACCENT_DARK = RGBColor(19, 66, 74)
MUTED = RGBColor(86, 101, 110)
LIGHT_FILL = "E9F4F5"
HEADER_FILL = "1D6F78"
QUOTE_FILL = "F2F7F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_no_wrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = OxmlElement("w:noWrap")
    tc_pr.append(no_wrap)


def clear_cell(cell):
    cell.text = ""


def add_run_with_inline_code(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = ACCENT_DARK
        else:
            paragraph.add_run(part)


def add_paragraph(doc, text, style=None, color=None, bold=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    add_run_with_inline_code(paragraph, text)
    for run in paragraph.runs:
        if color:
            run.font.color.rgb = color
        run.bold = bold or run.bold
    return paragraph


def normalize_table_row(line):
    raw = line.strip().strip("|")
    return [cell.strip() for cell in raw.split("|")]


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(table.columns)):
            cell = table.cell(r_idx, c_idx)
            clear_cell(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_run_with_inline_code(p, text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                set_cell_text_color(cell, RGBColor(255, 255, 255))
                for run in p.runs:
                    run.bold = True
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F7FBFB")
    doc.add_paragraph()


def apply_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(29, 35, 39)

    for name, size, color in [
        ("Heading 1", 20, ACCENT_DARK),
        ("Heading 2", 15, ACCENT),
        ("Heading 3", 12, ACCENT_DARK),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    quote = styles["Intense Quote"]
    quote.font.name = "Aptos"
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = ACCENT_DARK


def add_cover(doc, title, subtitle):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("PROTOGENESIS: PRIMORDIA")
    run.font.name = "Aptos Display"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Game Design Document v5.12")
    run.font.name = "Aptos Display"
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = MUTED

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, 220, 260, 220, 260)
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Documento de direccion final: juego terminado, cortes de scope y canon de producto.")
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK

    doc.add_page_break()


def build(markdown_path, output_path):
    text = Path(markdown_path).read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    apply_styles(doc)
    add_cover(doc, "PROTOGENESIS: PRIMORDIA", "Biological RTS de cuerpo multicelular, squads libres, Genome Lab y madre germinal latente.")

    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="No Spacing")
                p.paragraph_format.left_indent = Cm(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                for code_line in code_lines:
                    run = p.add_run(code_line + "\n")
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(45, 55, 60)
                in_code = False
                code_lines = []
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(" ")
            run.font.size = Pt(2)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_sep(lines[i]):
                    table_lines.append(normalize_table_row(lines[i]))
                i += 1
            add_table(doc, table_lines)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).strip()
            if level == 1:
                doc.add_heading(title, level=1)
            elif level == 2:
                doc.add_heading(title, level=2)
            else:
                doc.add_heading(title, level=3)
            i += 1
            continue

        if stripped.startswith(">"):
            quote = stripped.lstrip(">").strip()
            p = add_paragraph(doc, quote, style="Intense Quote")
            p.paragraph_format.left_indent = Cm(0.35)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_run_with_inline_code(p, stripped[2:].strip())
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_run_with_inline_code(p, numbered.group(1))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_gdd_docx.py input.md output.docx")
    build(sys.argv[1], sys.argv[2])
